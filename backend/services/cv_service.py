"""cv_service.py — All Groq-powered CV operations."""
import json, io, re
from groq import Groq
from config import get_settings

settings = get_settings()
client = Groq(api_key=settings.groq_api_key)
MODEL = "llama-3.3-70b-versatile"


def _chat_json(messages: list, max_tokens: int = 2000) -> dict:
    try:
        response = client.chat.completions.create(
            model=MODEL, max_tokens=max_tokens,
            response_format={"type": "json_object"},
            messages=messages,
        )
        text = response.choices[0].message.content
        # Strip markdown fences if model wraps response
        text = text.strip()
        if text.startswith("```"):
            text = re.sub(r"^```[a-z]*\n?", "", text)
            text = re.sub(r"\n?```$", "", text)
        return json.loads(text)
    except Exception as e:
        print(f"[Groq JSON] Error: {e}")
        raise


def _chat_text(prompt: str, max_tokens: int = 800) -> str:
    try:
        response = client.chat.completions.create(
            model=MODEL, max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"[Groq Text] Error: {e}")
        raise


def extract_cv_text(content: bytes, filename: str) -> str:
    """Enhanced CV text extraction with better formatting preservation."""
    lower = filename.lower()
    
    if lower.endswith(".pdf"):
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        
        for page in doc:
            # Extract text with layout preservation
            text = page.get_text("text")
            
            # Clean up common PDF artifacts
            text = text.replace('\x00', '')  # Remove null bytes
            text = text.replace('\uf0b7', '•')  # Fix bullet points
            text = text.replace('\u2022', '•')  # Normalize bullets
            text = text.replace('\u2013', '-')  # En dash to hyphen
            text = text.replace('\u2014', '-')  # Em dash to hyphen
            
            # Remove excessive whitespace while preserving structure
            lines = text.split('\n')
            cleaned_lines = []
            for line in lines:
                line = ' '.join(line.split())  # Normalize spaces
                if line.strip():  # Only keep non-empty lines
                    cleaned_lines.append(line)
            
            text_parts.append('\n'.join(cleaned_lines))
        
        doc.close()
        return '\n\n'.join(text_parts)
    
    if lower.endswith((".docx", ".doc")):
        from docx import Document
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Preserve bullet points
                if paragraph.style.name.startswith('List'):
                    text = '• ' + text if not text.startswith('•') else text
                text_parts.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return '\n'.join(text_parts)
    
    # Plain text files
    text = content.decode("utf-8", errors="ignore")
    
    # Clean up common issues
    text = text.replace('\r\n', '\n')  # Normalize line endings
    text = text.replace('\r', '\n')
    
    # Remove excessive blank lines (more than 2 consecutive)
    while '\n\n\n' in text:
        text = text.replace('\n\n\n', '\n\n')
    
    return text.strip()


async def analyze_cv(cv_text: str) -> dict:
    system = """You are a strict, honest senior career coach and ATS expert. Analyze only what is actually written in the CV.
This tool is used by professionals across ALL fields — Engineering, HR, Sales, Marketing, Finance, Operations, Design, Legal, Healthcare, Education, and more.

CRITICAL RULES:
- Only extract skills, experiences, and achievements that are EXPLICITLY stated in the CV. Do not invent or infer.
- strengths must be based on real, specific content from the CV — not generic platitudes like "strong communicator".
- critical_gaps must be realistic gaps for their field based on what is missing from the CV — do not fabricate skills.
- rewritten_bullets must use the actual bullet points from the CV as the "original" — do not make up experience.
- If a section of the CV is empty or missing, say so honestly rather than inventing content.

Return a JSON object with EXACTLY these keys — no extras, no omissions:

name (string)
email (string, or empty)
phone (string, or empty)
summary (string, 2 sentences based on the actual CV content and their specific field)
ats_score (integer 0-100, based on keyword density, formatting clarity, measurable achievements, section structure)
score_breakdown (object): { "keywords": int, "formatting": int, "experience": int, "education": int, "skills": int } — all 0-100
skills (array of strings, ALL skills extracted from CV — be thorough: include technical skills, tools, soft skills, domain expertise, certifications, languages. Adapt to the candidate's field.)
experience_years (integer)
education (string)
strengths (array of exactly 3 strings, specific to this person's actual CV content and field — NOT generic platitudes)
critical_gaps (array of EXACTLY 4 objects): EVERY object MUST have these exact keys:
  { "skill": "string — name of the missing skill or gap relevant to their field",
    "severity": "critical" | "moderate" | "minor",
    "reason": "string — 1 sentence explaining why this gap hurts job prospects in their field" }
  IMPORTANT: critical_gaps MUST always return exactly 4 objects. Even excellent CVs have gaps.
  DO NOT return strings, only objects. DO NOT omit any of the three keys.
rewritten_bullets (array of 5 objects): each MUST have:
  { "original": "exact bullet text from CV",
    "rewritten": "improved bullet: strong action verb + specific number/metric + business impact relevant to their field",
    "improvement": "one sentence explaining what changed" }
recommended_certs (array of 3 objects): each has:
  { "priority": 1|2|3, "name": string, "provider": string,
    "reason": "tailored to this person's specific skill gaps and field",
    "score_impact": string like "+8 ATS points" }
recommended_projects (array of 3 objects): each has:
  { "title": string, "difficulty": "beginner"|"intermediate"|"advanced",
    "description": "2 sentences on what to build/do and why it helps in their field",
    "skills_added": [array of strings] }"""

    return _chat_json([
        {"role": "system", "content": system},
        {"role": "user", "content": f"Analyze this CV:\n\n{cv_text[:4000]}"},
    ], max_tokens=3000)


async def score_cv_against_jd(cv_text: str, jd_text: str) -> dict:
    system = """You are a strict, honest recruitment expert. Analyze the CV against the job description.

CRITICAL RULES — follow exactly:
- Only extract keywords that ACTUALLY appear in the job description. Do not invent keywords.
- Only list matched_keywords that genuinely appear in BOTH the CV and the JD. Do not hallucinate matches.
- missing_keywords must only include terms explicitly mentioned in the JD that are absent from the CV.
- tailored_bullets must reference real experience from the CV reworded to match the JD — do not invent projects or achievements.
- If the job description is vague, short, or nonsensical, set match_score to 0, verdict to "Invalid JD", and set application_strategy to "The job description provided does not contain enough information to perform a meaningful analysis. Please paste the full job description."
- Do NOT fabricate an analysis for a poor-quality JD. Be honest.
- application_strategy must name specific skills and experiences from THIS candidate's actual CV.

Return a JSON object with EXACTLY these keys:
match_score (integer 0-100)
verdict (string, one of: "Strong Match", "Good Match", "Partial Match", "Weak Match", "Invalid JD")
application_strategy (string, 2-3 sentences citing specific CV content and specific JD requirements)
matched_keywords (array of strings — only real overlaps between CV and JD)
missing_keywords (array of strings — only terms explicitly in the JD, absent from CV)
tailored_bullets (array of 3 strings — real CV experience reworded for this JD, no invented content)
rewritten_cv_sections (object):
  { "summary": "rewritten summary using only real CV content targeting this JD",
    "skills_to_highlight": [strings from their actual CV that match this JD],
    "skills_to_add": [strings explicitly required by JD that candidate lacks] }
recommendation (string, 2-3 honest, specific, actionable sentences)"""

    return _chat_json([
        {"role": "system", "content": system},
        {"role": "user", "content": f"CV:\n{cv_text[:2500]}\n\nJob Description:\n{jd_text[:2000]}"},
    ], max_tokens=1500)


async def generate_full_cover_letter(cv_text: str, jd_text: str, company: str, role: str) -> str:
    prompt = f"""Write a professional, personalized cover letter for a candidate applying to {role} at {company}.
3-4 paragraphs. No placeholder text. Reference specific projects and skills from their CV.
Return only the letter text, nothing else.

CV:
{cv_text[:2500]}

Job Description:
{jd_text[:1500]}"""
    return _chat_text(prompt, max_tokens=900)


async def generate_interview_prep(cv_text: str, role: str, company: str = "") -> list:
    company_str = f" at {company}" if company else ""
    system = f"""You are a senior interviewer for {role} roles{company_str}.
Return a JSON object with key "questions" containing an array of 8 objects.
Each object: {{ "question": string, "category": string (one of: technical, behavioural, project, system-design), "tip": string }}
Base questions specifically on the candidate's actual CV projects and skills."""

    data = _chat_json([
        {"role": "system", "content": system},
        {"role": "user", "content": f"Generate interview questions based on:\n\n{cv_text[:2500]}"},
    ], max_tokens=1500)
    return data.get("questions", []) if isinstance(data, dict) else data


async def evaluate_practice_answer(cv_text: str, question: str, user_answer: str, role: str) -> dict:
    system = """You are a strict but fair interview coach. Evaluate the candidate's answer.
Return a JSON object with EXACTLY:
score (integer 1-10)
verdict (string, one of: "Excellent", "Good", "Needs Work", "Too Vague")
what_worked (array of 2-3 strings)
what_to_improve (array of 2-3 strings)
ideal_answer_outline (string, 2-3 sentences describing what a 10/10 answer looks like)"""

    return _chat_json([
        {"role": "system", "content": system},
        {"role": "user", "content": f"Role: {role}\nQuestion: {question}\nCandidate's answer: {user_answer}\n\nCV context:\n{cv_text[:1500]}"},
    ], max_tokens=800)
